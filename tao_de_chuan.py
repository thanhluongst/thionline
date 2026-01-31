from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def create_advanced_sample_docx():
    doc = Document()
    
    # Tiêu đề
    title = doc.add_heading('ĐỀ THI MẪU CHUẨN ĐỊNH DẠNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- PHẦN 1: TRẮC NGHIỆM 4 PHƯƠNG ÁN ---
    doc.add_heading('Phần I. Câu trắc nghiệm nhiều lựa chọn', level=1)
    doc.add_paragraph('Mỗi câu hỏi thí sinh chỉ chọn một phương án.')

    # Câu 1
    doc.add_paragraph('Câu 1. Tập xác định của hàm số y = log2(x - 3) là:')
    doc.add_paragraph('A. D = (3; +∞)')
    # Đáp án đúng B gạch chân
    p1b = doc.add_paragraph('B. ')
    run_1b = p1b.add_run('D = [3; +∞)')
    run_1b.underline = True 
    doc.add_paragraph('C. D = R')
    doc.add_paragraph('D. D = (0; +∞)')

    # Câu 2
    doc.add_paragraph('Câu 2. Nguyên hàm của hàm số f(x) = cosx là:')
    doc.add_paragraph('A. sinx + C')
    doc.add_paragraph('B. -sinx + C')
    doc.add_paragraph('C. tanx + C')
    run_2d = doc.add_paragraph('D. ').add_run('cosx + C')
    run_2d.underline = True

    # --- PHẦN 2: ĐÚNG SAI (Mỗi câu có 4 ý a, b, c, d) ---
    doc.add_heading('Phần II. Câu trắc nghiệm đúng sai', level=1)
    doc.add_paragraph('Câu 3. Cho hàm số y = ax^3 + bx^2 + cx + d có đồ thị như hình vẽ.')
    
    run_3a = doc.add_paragraph('a) ').add_run('Hàm số đồng biến trên khoảng (0; 2).')
    run_3a.underline = True # Đúng
    
    doc.add_paragraph('b) Đồ thị hàm số có hai điểm cực trị.')
    
    run_3c = doc.add_paragraph('c) ').add_run('Giá trị cực đại của hàm số bằng 4.')
    run_3c.underline = True # Đúng
    
    doc.add_paragraph('d) Phương trình f(x) = 0 có 3 nghiệm phân biệt.')

    # --- PHẦN 3: ĐIỀN KHUYẾT ---
    doc.add_heading('Phần III. Câu trắc nghiệm trả lời ngắn (Điền khuyết)', level=1)
    
    doc.add_paragraph('Câu 4. Cho hình chóp S.ABCD có đáy là hình vuông cạnh a, SA vuông góc với đáy. Biết SA = a√3. Tính góc giữa đường thẳng SC và mặt phẳng (ABCD) (đơn vị: độ).')
    run_4_ans = doc.add_paragraph('Đáp án: ').add_run('60')
    run_4_ans.underline = True

    # Lưu file
    doc.save('De_Thi_Chuan.docx')
    print("--- Đã tạo file De_Thi_Chuan.docx ---")
    print("Mẹo: Mở file bằng Word để kiểm tra các dòng gạch chân.")

if __name__ == "__main__":
    create_advanced_sample_docx()
